"""
晓牧传媒 · 核心生成逻辑
路径全部相对于本文件，可部署到任何环境
"""
import openai, json, re, os, datetime, io, logging

logger = logging.getLogger(__name__)
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROMPTS_DIR = Path(__file__).parent / 'prompts'

# ── 工具函数 ──────────────────────────────────────

def load_prompt(filename):
    return (PROMPTS_DIR / filename).read_text(encoding='utf-8')

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def difficulty_style(d):
    if d == '简单版': return (46,125,50),  'E8F5E9'
    if d == '标准版': return (230,81,0),   'FFF3E0'
    return                  (106,27,154),  'EDE7F6'

# ── 口吻规范 ─────────────────────────────────────

# ── 违禁词库（程序级扫描）────────────────────────────
_FW_PATH = PROMPTS_DIR / 'forbidden_words.json'
_FW_DATA  = json.loads(_FW_PATH.read_text(encoding='utf-8')) if _FW_PATH.exists() else {}
_FW_ROUTING = _FW_DATA.get('_routing', {})

def _get_industry_keys(client_text: str) -> list:
    """根据客户信息文本，返回匹配的违禁词分类key列表"""
    keys = []
    for kw, cat in _FW_ROUTING.items():
        if kw in client_text and cat not in keys:
            keys.append(cat)
    return keys

# 禁止开场词（程序级检测，触发自动重写）
FORBIDDEN_OPENING_PATTERNS = [
    "大家好", "你好，我是", "欢迎大家", "今天给大家",
    "欢迎来到", "我来给大家", "今天我来", "我在这里",
]

# 场景禁止出现的人物词（C2/C3规则）
SCENE_FORBIDDEN_PERSONS = [
    "团队成员", "设计师", "员工", "顾客", "客户", "年轻时", "同事",
    "助理", "工人", "师傅们", "我们团队",
]

# ── 程序级后处理：「最」字替换 ───────────────────────
# 替换规则：最X → 替代表达，直接字符串替换，不依赖AI

_ZUI_REPLACEMENTS = [
    # 两字组合（优先匹配长词）
    ("最好吃", "好吃"),
    ("最好喝", "好喝"),
    ("最好的", "很好的"),   # 长词优先，避免被"最好"截断
    ("最新鲜", "非常新鲜"),
    ("最实惠", "划算"),
    ("最便宜", "价格低"),
    ("最贵",   "价格高"),
    ("最香",   "很香"),
    ("最辣",   "够辣"),
    ("最甜",   "很甜"),
    ("最快",   "够快"),
    ("最慢",   "比较慢"),
    ("最好",   "挺好"),
    ("最差",   "比较差"),
    ("最难",   "确实难"),
    ("最简单", "比较简单"),
    ("最重要", "很关键"),
    ("最关键", "关键"),
    ("最核心", "核心"),
    ("最大的", "很大"),
    ("最小的", "很小"),
    ("最多",   "数量多"),
    ("最少",   "数量少"),
    ("最高",   "很高"),
    ("最低",   "很低"),
    ("最强",   "很强"),
    ("最弱",   "比较弱"),
    ("最专业", "够专业"),
    ("最厉害", "很厉害"),
    ("最靠谱", "靠谱"),
    ("最真诚", "真诚"),
    ("最透明", "透明"),
    ("最放心", "放心"),
    ("最安全", "安全"),
    ("最稳",   "稳"),
    ("最优质", "品质好"),
    ("最高端", "比较高端"),
    ("最高级", "比较高级"),
    ("最受欢迎", "受欢迎"),
    ("最受",   "很受"),
    ("最爱",   "喜欢"),
    ("最喜欢", "挺喜欢"),
    ("最满意", "很满意"),
    ("最担心", "很担心"),
    ("最正宗", "很地道"),
    ("最地道", "很地道"),
    ("最有效", "非常有效"),
    ("最真实", "很真实"),
    ("最宝贵", "很宝贵"),
    ("最鲜活", "非常鲜活"),
    ("最难忘", "很难忘"),
    ("最怕",   "最怕"),      # 保留：表达恐惧，不是吹嘘
    ("最后",   "最后"),      # 保留：时间序列用词
    ("最终",   "最终"),      # 保留：结果用词
    ("最近",   "最近"),      # 保留：时间用词
    ("最初",   "最初"),      # 保留：时间用词
]

# ── D1 极限词替换表（绝对/第一/独一无二 等，长词优先）──
_D1_REPLACEMENTS = [
    ("行业领先者",    "行业前列"),
    ("行业第一",      "行业前列"),
    ("销量第一",      "销量领先"),
    ("全国第一",      "全国领先"),
    ("品质第一",      "品质上乘"),
    ("绝对值得",      "真的值得"),
    ("绝对好吃",      "真的好吃"),
    ("绝对让您满意",  "保证让您满意"),
    ("绝对好",        "真的好"),
    ("绝对",          "确实"),
    ("史上最强",      "非常强"),
    ("独一无二",      "独特"),
    ("无与伦比",      "出色"),
    ("无可替代",      "不可或缺"),
]

# 需要保留的「最」（不替换）
_ZUI_KEEP = {"最后", "最终", "最近", "最初", "最怕"}

# ── 词频限制表（全批合计超出次数则替换后续出现）──────
_FREQ_LIMITS = {
    '每一个': (3, '每个'),
    '每一件': (3, '每件'),
    '每一位': (3, '每位'),
    '每一道': (3, '每道'),
    '每一处': (3, '每处'),
    '每一次': (3, '每次'),
    '匠心':   (2, '认真'),
    '追求':   (3, '要求'),
    '坚持':   (4, '一直'),
    '品质':   (5, '质量'),
    '用心':   (3, '认真'),
    '初心':   (2, '当初的想法'),
    '情怀':   (1, '感情'),
    '不仅仅是': (3, '不只是'),
    '更是':    (3, '也是'),
    '每一台':  (3, '每台'),
    '致力于': (2, '一直在做'),
    '专注于': (2, '专门做'),
    '赋能':   (1, '帮助'),
    '全程无忧': (1, '放心交给我'),
    '一站式': (1, '一条龙'),
}

# ── 结尾模式判定 ─────────────────────────────────────
_HARD_ENDING_PATTERNS = [
    '欢迎来', '私信我', '来店', '来尝', '来找我',
    '来工作室', '来厂', '联系我', '扫码', '预约',
    '点击主页', '主页预约', '来这里', '等你来',
]
_SOFT_ENDING_PATTERNS = [
    '你觉得', '你遇到过', '你会选', '你怎么看',
    '评论告诉', '说说你', '你有没有', '你呢',
    '你有过这种', '你们那边', '你家',
]

def _classify_ending(dialogue: str) -> str:
    for p in _HARD_ENDING_PATTERNS:
        if p in dialogue:
            return 'hard'
    for p in _SOFT_ENDING_PATTERNS:
        if p in dialogue:
            return 'soft'
    return 'natural'

def _get_last_dialogue(script) -> str:
    shots = script.get('shots', [])
    for shot in reversed(shots):
        d = shot.get('dialogue', '').strip()
        if d:
            return d
    return ''

def _clear_last_hard_ending(script):
    """把最后一个镜头的硬引导词删掉，改为自然收尾"""
    shots = script.get('shots', [])
    for i in range(len(shots) - 1, -1, -1):
        d = shots[i].get('dialogue', '')
        if any(p in d for p in _HARD_ENDING_PATTERNS):
            for p in _HARD_ENDING_PATTERNS:
                d = d.replace(p, '')
            # 清理多余标点
            d = re.sub(r'[，。！、\s]+$', '。', d.strip())
            shots[i]['dialogue'] = d
            return

_post_process_call_count = 0

def post_process_scripts(scripts, client_name=''):
    """程序级后处理：最字替换 + 词频限制 + 结尾配额 + C2场景修正"""
    global _post_process_call_count
    _post_process_call_count += 1
    call_n = _post_process_call_count
    logger.info("[post_process #%d] 开始处理，共 %d 条脚本，client=%r", call_n, len(scripts), client_name)
    replacements_made = []

    # ── Step 1: 极限词替换（「最」字 + D1 极限词）────────
    for s in scripts:
        for field_name, value in [('title', s.get('title', '')),
                                   ('tips',  s.get('tips',  ''))]:
            new_val = _apply_d1_replacements(_apply_zui_replacements(value))
            if new_val != value:
                replacements_made.append({'script': s['number'], 'field': field_name, 'type': 'zui'})
                s[field_name] = new_val
        for shot in s.get('shots', []):
            for field in ('dialogue', 'scene'):
                original = shot.get(field, '')
                replaced = _apply_d1_replacements(_apply_zui_replacements(original))
                if replaced != original:
                    replacements_made.append({'script': s['number'], 'field': field, 'type': 'zui'})
                    shot[field] = replaced

    # ── Step 2: 词频限制（全批统计超出则替换）──────────
    # 收集全部口播+标题文本（按顺序）
    word_counters = {w: 0 for w in _FREQ_LIMITS}
    for s in scripts:
        full_parts = [s.get('title', '')] + [
            sh.get('dialogue', '') for sh in s.get('shots', [])
        ]
        for part_idx, part in enumerate(full_parts):
            new_part = part
            for word, (limit, replacement) in _FREQ_LIMITS.items():
                # 先统计本段内该词出现次数，逐个决定是否替换
                occurrences = new_part.count(word)
                for _ in range(occurrences):
                    word_counters[word] += 1
                    if word_counters[word] > limit:
                        new_part = new_part.replace(word, replacement, 1)
                        replacements_made.append({'script': s['number'], 'field': 'freq', 'type': word})
            if new_part != part:
                if part_idx == 0:
                    s['title'] = new_part
                else:
                    s['shots'][part_idx - 1]['dialogue'] = new_part

    # ── Step 3: 场景描述修正（C2多人出镜 + C4回忆复刻）──
    name = client_name or '主理人'
    c2_fix_count = 0
    for s in scripts:
        for shot in s.get('shots', []):
            scene = shot.get('scene', '')
            for kw in SCENE_FORBIDDEN_PERSONS:
                if kw in scene:
                    shot['scene'] = re.sub(
                        r'.{0,8}' + re.escape(kw) + r'.{0,8}',
                        f'{name}站在镜头前',
                        scene
                    )
                    c2_fix_count += 1
                    replacements_made.append({'script': s['number'], 'field': 'scene_c2', 'type': kw})
                    break
    for s in scripts:
        for shot in s.get('shots', []):
            scene = shot.get('scene', '')
            for kw in C4_SCENE_KEYWORDS:
                if kw in scene:
                    shot['scene'] = f'{name}站在当前位置，讲述过去的经历，镜头对着他'
                    replacements_made.append({'script': s['number'], 'field': 'scene_c4', 'type': kw})
                    break

    # ── Step 4: 结尾配额控制（硬引导≤6条）───────────
    hard_count = 0
    HARD_LIMIT = 6
    for s in scripts:
        last_d = _get_last_dialogue(s)
        ending_type = _classify_ending(last_d)
        s['_ending_type'] = ending_type
        if ending_type == 'hard':
            hard_count += 1
            if hard_count > HARD_LIMIT:
                _clear_last_hard_ending(s)
                s['_ending_type'] = 'natural'
                replacements_made.append({'script': s['number'], 'field': 'ending_quota', 'type': 'hard→natural'})

    # ── Step 5: 重复检测（同批内相似度 >60% 的脚本打标）──
    # 提取每条脚本的全部口播句子集合，用于两两比较
    def _dialogue_sentences(s):
        sents = set()
        for shot in s.get('shots', []):
            d = shot.get('dialogue', '')
            # 按标点分句
            for part in re.split(r'[，。！？、……——\n]+', d):
                part = part.strip()
                if len(part) >= 5:
                    sents.add(part)
        return sents

    seen_sentences = []   # list of (script_number, set_of_sentences)
    for s in scripts:
        curr = _dialogue_sentences(s)
        if not curr:
            seen_sentences.append((s['number'], curr))
            continue
        for prev_num, prev_sents in seen_sentences:
            if not prev_sents:
                continue
            overlap = len(curr & prev_sents)
            ratio = overlap / min(len(curr), len(prev_sents))
            if ratio >= 0.6:
                replacements_made.append({
                    'script': s['number'],
                    'field': 'dedup_flag',
                    'type': f'与第{prev_num}条重复率{int(ratio*100)}%',
                })
                break
        seen_sentences.append((s['number'], curr))

    zui_n   = sum(1 for r in replacements_made if r.get('type') == 'zui')
    freq_n  = sum(1 for r in replacements_made if r.get('field') == 'freq')
    c2_n    = sum(1 for r in replacements_made if r.get('field') == 'scene_c2')
    c4_n    = sum(1 for r in replacements_made if r.get('field') == 'scene_c4')
    dedup_n = sum(1 for r in replacements_made if r.get('field') == 'dedup_flag')
    end_n   = sum(1 for r in replacements_made if r.get('field') == 'ending_quota')
    logger.info(
        "[post_process #%d] 完成。总替换 %d 处：极限词=%d, 词频=%d, C2场景=%d, C4回忆=%d, 结尾配额=%d, 重复标记=%d",
        call_n, len(replacements_made), zui_n, freq_n, c2_n, c4_n, end_n, dedup_n,
    )
    return scripts, replacements_made


def _apply_zui_replacements(text):
    """对单段文本执行「最」字替换，保留白名单词"""
    if '最' not in text:
        return text
    for keep in _ZUI_KEEP:
        # 临时保护白名单词：用占位符替换
        text = text.replace(keep, f'__KEEP_{keep}__')
    for old, new in _ZUI_REPLACEMENTS:
        text = text.replace(old, new)
    # 清理残余的孤立「最」（前后都是中文字符）
    text = re.sub(r'(?<=[\u4e00-\u9fff])最(?=[\u4e00-\u9fff])', '', text)
    # 还原白名单词
    for keep in _ZUI_KEEP:
        text = text.replace(f'__KEEP_{keep}__', keep)
    return text


def _apply_d1_replacements(text):
    """对单段文本执行D1极限词替换（绝对/行业第一/独一无二等）"""
    for old, new in _D1_REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
    return text

def scan_forbidden_words(scripts, prompt_file, industry_keys=None):
    """返回违规列表 [{'script': 条号, 'word': 违禁词}]"""
    # 通用词
    words = list(_FW_DATA.get('all', []))
    # 按行业路由加载专项词
    for key in (industry_keys or []):
        words += _FW_DATA.get(key, [])
    # 旧有 prompt_file 映射（兼容）
    legacy = {
        'jiazhuang': ["无甲醛", "零甲醛", "无毒"],
        'canyin':    ["纯天然", "无添加", "祖传秘方"],
    }
    words += legacy.get(prompt_file.replace('.md', ''), [])

    violations = []
    for s in scripts:
        full_text = s.get('title', '') + s.get('tips', '') + ''.join(
            shot.get('dialogue', '') + shot.get('scene', '')
            for shot in s.get('shots', [])
        )
        for w in words:
            if w in full_text:
                violations.append({'script': s.get('number', '?'), 'word': w})
    return violations

def scan_forbidden_openings(scripts):
    """检测开场套话，返回违规条号列表"""
    violations = []
    for s in scripts:
        shots = s.get('shots', [])
        if not shots:
            continue
        first_line = shots[0].get('dialogue', '')
        first_12 = first_line[:12]
        hit = any(p in first_12 for p in FORBIDDEN_OPENING_PATTERNS)
        # 也检测 "我是老X / 我是小X" 格式
        if not hit and re.match(r'^我是[\u4e00-\u9fff]{1,4}', first_line):
            hit = True
        if hit:
            violations.append(s['number'])
    return violations

def scan_scene_persons(scripts):
    """检测场景描述中出现的禁止人物词，返回违规详情"""
    violations = []
    for s in scripts:
        for i, shot in enumerate(s.get('shots', [])):
            scene = shot.get('scene', '')
            for kw in SCENE_FORBIDDEN_PERSONS:
                if kw in scene:
                    violations.append({
                        'script': s.get('number', '?'),
                        'shot': i + 1,
                        'keyword': kw,
                    })
                    break
    return violations

# ── 自动评分辅助 ─────────────────────────────────────

_D1_SCORE_WORDS = [
    '唯一', '绝对', '独一无二', '无与伦比', '无可替代',
    '行业第一', '销量第一', '全国第一', '史上最强',
]

def _count_remaining_extreme_words(scripts):
    """统计D1极限词在postProcess后的残余数量"""
    count = 0
    for s in scripts:
        text = ' '.join(
            [s.get('title', ''), s.get('tips', '')]
            + [shot.get('dialogue', '') + shot.get('scene', '')
               for shot in s.get('shots', [])]
        )
        for w in _D1_SCORE_WORDS:
            count += text.count(w)
        for keep in _ZUI_KEEP:
            text = text.replace(keep, '')
        count += text.count('最')
    return count

def _count_long_sentences(scripts, threshold=25):
    """统计口播中超过 threshold 字的单句数量"""
    count = 0
    for s in scripts:
        for shot in s.get('shots', []):
            for sent in re.split(r'[，。！？、……——\n]+', shot.get('dialogue', '')):
                if len(sent.strip()) > threshold:
                    count += 1
    return count

def _endings_count(scripts):
    """统计三类结尾的数量"""
    result = {'hard': 0, 'soft': 0, 'natural': 0}
    for s in scripts:
        t = _classify_ending(_get_last_dialogue(s))
        result[t] = result.get(t, 0) + 1
    return result

def _dup_pair_count(scripts):
    """统计重复对数（口播句子集合相似度≥60%）"""
    def _sents(s):
        sents = set()
        for shot in s.get('shots', []):
            for part in re.split(r'[，。！？、……——\n]+', shot.get('dialogue', '')):
                part = part.strip()
                if len(part) >= 5:
                    sents.add(part)
        return sents
    pairs, seen = 0, []
    for s in scripts:
        curr = _sents(s)
        if curr:
            for prev in seen:
                if prev and len(curr & prev) / min(len(curr), len(prev)) >= 0.6:
                    pairs += 1
                    break
        seen.append(curr)
    return pairs

def auto_score(scripts, client):
    """
    五维评分 + 合规扣分，满分100分，交付线90分
    返回 {'total': int, 'breakdown': dict, 'issues': list}
    """
    bd, issues = {}, []
    prompt_file   = client.get('prompt_file', 'general.md')
    industry_keys = client.get('industry_keys', [])
    long_n        = _count_long_sentences(scripts, 25)

    # ① 钩子力 25分
    bad_hooks = scan_forbidden_openings(scripts)
    penalty = len(bad_hooks) * 3
    bd['钩子力'] = max(0, 25 - penalty)
    if bad_hooks:
        issues.append(f"B1套话开场：第{'、'.join(str(n) for n in bad_hooks)}条（-{penalty}分）")

    # ② 内容自然感 20分
    all_dialogue = ' '.join(
        shot.get('dialogue', '')
        for s in scripts for shot in s.get('shots', [])
    )
    nat_penalty = 0
    for word, (limit, _) in _FREQ_LIMITS.items():
        if all_dialogue.count(word) > limit * 2:
            nat_penalty += 2
    nat_penalty += min(4, long_n // 30)
    bd['内容自然感'] = max(8, 20 - nat_penalty)

    # ③ 内容多样性 20分
    type_counts  = {}
    for s in scripts:
        t = s.get('type', '其他')
        type_counts[t] = type_counts.get(t, 0) + 1
    interactive_n = type_counts.get('互动型', 0)
    dup_pairs     = _dup_pair_count(scripts)
    div_penalty   = max(0, 5 - interactive_n) * 2 + dup_pairs * 3
    bd['内容多样性'] = max(10, 20 - div_penalty)
    if interactive_n < 5:
        issues.append(f"互动型偏少：{interactive_n}条（参考值≥5条）")
    if dup_pairs:
        issues.append(f"疑似重复对：{dup_pairs}对（-{dup_pairs * 3}分）")

    # ④ 拍摄可执行性 18分
    c2_n = len(scan_scene_persons(scripts))
    c4_n = len(scan_c4_scenes(scripts))
    shoot_penalty = (c2_n // 5) * 2 + c4_n * 2
    bd['拍摄可执行性'] = max(8, 18 - shoot_penalty)
    if c2_n:
        issues.append(f"C2多人出镜：{c2_n}处（-{(c2_n // 5) * 2}分）")
    if c4_n:
        issues.append(f"C4回忆复刻：{c4_n}处（-{c4_n * 2}分）")

    # ⑤ 结尾节奏 17分
    endings      = _endings_count(scripts)
    hard_penalty = max(0, endings['hard'] - 6)
    soft_penalty = int(max(0, 8 - endings['soft']) * 0.5)
    bd['结尾节奏'] = max(6, 17 - hard_penalty - soft_penalty)
    if endings['hard'] > 6:
        issues.append(f"硬引导偏多：{endings['hard']}条（上限6条，-{hard_penalty}分）")

    # D1 极限词
    d1_n = _count_remaining_extreme_words(scripts)
    bd['D1极限词'] = -min(8, d1_n)
    if d1_n:
        issues.append(f"D1极限词残余：{d1_n}次（-{min(8, d1_n)}分）")

    # D2 行业违禁词
    d2_violations = scan_forbidden_words(scripts, prompt_file, industry_keys)
    d2_n = len(d2_violations)
    bd['D2违禁词'] = -min(15, d2_n * 5)
    if d2_n:
        issues.append(f"D2行业违禁词：{d2_n}个（-{min(15, d2_n * 5)}分）")

    # D3 长句
    bd['D3长句'] = -min(5, int(long_n * 0.5))

    total = max(0, sum(bd.values()))
    label = "✅达标" if total >= 90 else ("⚠️良好" if total >= 80 else "❌偏低")
    logger.info(
        "[auto_score] %s client=%r 总分=%d | 钩子=%d 自然感=%d 多样性=%d 可执行=%d 结尾=%d D1=%d D2=%d D3=%d",
        label, client.get('name', '?'), total,
        bd['钩子力'], bd['内容自然感'], bd['内容多样性'],
        bd['拍摄可执行性'], bd['结尾节奏'], bd['D1极限词'], bd['D2违禁词'], bd['D3长句'],
    )
    return {'total': total, 'breakdown': bd, 'issues': issues}

C4_SCENE_KEYWORDS = ['回忆镜头', '老照片', '旧照片', '当年的照片', '年轻时的', '历史画面', '闪回']

def scan_c4_scenes(scripts):
    """检测场景描述中出现的C4回忆复刻关键词"""
    violations = []
    for s in scripts:
        for i, shot in enumerate(s.get('shots', [])):
            scene = shot.get('scene', '')
            for kw in C4_SCENE_KEYWORDS:
                if kw in scene:
                    violations.append({'script': s.get('number', '?'), 'shot': i + 1, 'keyword': kw})
                    break
    return violations

def extract_city(location_text):
    """从完整地址中只提取城市/商圈名"""
    if not location_text:
        return location_text
    # 去掉路/号/大厦/楼层等详细信息
    clean = re.sub(r'(路|街|道|大道|大街)[\d零一二三四五六七八九十百\s]*号.*', '', location_text)
    clean = re.sub(r'(大厦|广场|中心|楼|层|室|号楼|栋|幢|单元).*', '', clean)
    clean = clean.strip()
    return clean if clean else location_text

# ── 规则包 v2.0（每批5条前重注入，防止规则衰减）───────
RULES_V2 = """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【规则包 v2.0 — 本批生成前必须完整读取，每条脚本均适用】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

■ R1【开场钩子 — 5类钩子必须覆盖全组30条，每类至少4条】
⚠️ 套话开场（大家好/我是XX/欢迎/我在这里）程序自动扫描，发现即退回重写。

【5类钩子必须分布在30条里，禁止某类一条都没有】

🔴 恐惧式 — 让人产生"我会不会也有这个问题"的担忧（至少4条）
  写法：揭露危险/错误/代价，让人感到紧迫
  示例："如果你用的护肤品3天就见效，快停下！这背后可能有你不知道的东西。"
  示例："很多人以为省了这笔钱，结果后面返工花的是双倍——我见过太多这样的单子。"

🟡 焦虑式 — 用数据/现象制造紧迫感，让人觉得"再不看就晚了"（至少4条）
  写法：数字+反常识，或行业现象+尖锐结论
  示例："做了10年皮肤调理，我发现80%的人都在犯同一个错误——"
  示例："同一条街上开了5家同行，4家都关了。不是竞争太激烈，是有一件事他们都没做对。"

🟠 借势式 — 借行业话题/趋势/公众事件引入，体现专业判断（至少4条）
  写法：最近发生了什么 + 主理人有自己的看法/立场
  示例："最近网上都在说这款成分，作为从业者，我有话要说——"
  示例："这个行业今年变化很大，有件事让我感触很深——"

🔵 好奇式 — 用反常识结论、悬念、未揭晓的谜题，让人想继续看（至少4条）
  写法：抛出反直觉的现象或问题，故意不给答案
  示例："同样的问题，有人花5万没解决，有人花200块就好了——区别到底在哪儿？"
  示例："5年前，我就站在这间空屋子里，所有人都说我疯了。（停顿）但今天，我想给所有不信我的人看看……"

🟢 共情式 — 描述受众共同经历的痛苦/困境，让人感到"这说的就是我"（至少4条）
  写法：你有没有/你是不是 + 具体痛苦场景，再引出主理人的理解
  示例："你是不是也试过很多方法，钱花了不少，效果却越来越失望——我懂这种感受。"
  示例："有客户跟我说，装修这件事，比结婚还累。我听完，真的笑不出来。"

开场自检标准：①首句≤22字 ②读出来自然有停顿 ③听完第一句想继续听。三条都要达到。

■ R2【出镜限制】只有主理人一人出镜。场景描述禁止出现团队/设计师/员工/顾客出镜（程序会扫描场景描述关键词）。客户故事只在口播中说，镜头只拍主理人。

■ R3【场景可执行】限门店/工作室内部。一部手机可完成为标准。不得要求：无人机/轨道推车/多机位/专业打光/数控设备界面特写。

■ R4【回忆处理】涉及过去的场景只通过主理人口播讲述。不得要求镜头还原历史/回忆画面。正确写法：「主理人站在工作室，讲述过去的经历，镜头对着他」

■ R5【口语节奏】句子长短自然交替，不要每句都一样长。
  短句示例（1-8字）："真的头大。" / "就这一件事。" / "没想到——"
  长句示例（20-35字）："那天我跟客户说，这个柜子我可以做，但你要给我三天时间，我不赶工。"
  禁止：所有句子都在15-20字左右（这是AI味的来源）。禁止书面排比句和连续四字成语。
  每条脚本总口播150-200字，保证内容饱满但不拖沓。

■ R6【违禁词】全行业禁：无与伦比/史上最强/行业领先/行业第一/专业团队/匠心/初心。
  建材额外禁：无甲醛/零甲醛/无毒。食品额外禁：纯天然/无添加/祖传秘方。
  医疗额外禁：根治/治愈/无副作用。教育额外禁：保过/保分/承诺成绩。

■ R7【结尾引导 + 金句收尾】禁套话：等你/理想家园/一起打造/欢迎关注/我在这里等你/点个关注。
  结尾策略按内容类型区分：
  · 故事型/幕后型/教育型/观点型/氛围型/对比型 → 自然收尾，末句用一句有力的金句点题，不加任何引导
  · 痛点型 → 末句给主理人自己的做法，点到为止，可用金句收
  · 互动型 → 可以有一句轻量互动问句（"你遇到过这种情况吗"）
  · 产品型 → 最多一句轻量提示，禁止强推
  整组30条：有引导结尾的不超过8条，其余全部自然收尾。

  【金句收尾标准 — 自然收尾类型的末句必须达到这个水平】
  金句 = 有态度 + 有情绪 + 简短有力，让人听完想截图保存的那种
  ✓ "5年，我没学会怎么快，但我学会了不将就。"
  ✓ "我不是在做生意，我是在证明一件事——认真的人，不会输。"
  ✓ "这行水深，但只要你足够真，就没人能把你淹死。"
  ✓ "后来我明白了，不是客户挑剔，是我当初没有底线。"
  ✓ "做到现在，我最怕的不是差评，是做出来的东西连自己都不信。"
  × 禁止："感谢大家的支持，我会继续努力的。"（鸡汤套话）
  × 禁止："希望我的经历对你有帮助。"（收尾套话）
  × 禁止："我会一直坚持下去的。"（空洞，无态度）
  金句必须来自脚本本身的主题/事件，不能贴标签式强加。

■ R8【地址】禁止完整街道/门牌号/楼层号。只允许城市名或知名商圈名。

■ R9【写作深度 — 区分好稿和烂稿的核心标准】
  每条脚本必须做到：
  ① 有一个具体细节：具体年份 / 具体金额 / 具体事件名 / 具体地名，不接受模糊概括
  ② 有情绪张力或冲突：让人听完第一句就想听下一句，不能全程平铺陈述
  ③ 口语自然有克制：说话像真人，有停顿（……/——），不说连续排比句

  【好稿写法示例 — 每条都要达到这个水平】
  ✓ "做了这行十几年，有一件事我从没在镜头前说过……"（悬念+confession，让人继续看）
  ✓ "2009年进这行，第一份工作月薪八百，每天站十二个小时。"（具体数字+代价感）
  ✓ "跟你聊一单，到现在我还记得——做起来真的头大。"（个人叙事+情绪锚）
  ✓ "有个事，我做了十几年之后才发现，当时真的想不明白。"（反转+经验感）
  ✓ "这个城市有多少人装修时吃过亏，我大概知道。"（地方感+权威感）
  ✓ "有人问我，你们家和别人家到底差在哪儿——我说，你先让我给你看一样东西。"（互动+引导）

  【不允许的写法 — 出现即视为低质量】
  × "我们的工艺一直追求高标准，客户都非常满意。"（概括评价，无细节）
  × "多年来我坚持用心做好每一件产品。"（空洞，无具体事件）
  × "我们的服务让每一位客户感到满意。"（宣传腔，无真实感）
  × "这个行业水很深，大家一定要注意选择。"（泛泛警告，无干货）

  写完每条脚本，对照这4条烂稿特征自检一遍。有一条中标就需要重写。

■ R10【第一人称 — 这条违反等于整条作废】
  所有口播台词必须是主理人本人用"我"开口讲话。
  ✓ 正确："我当时真不知道该怎么办——"
  ✓ 正确："我进这行的时候，口袋里只剩三百块。"
  × 错误："川哥常说，做这行要用心。"（第三人称旁白）
  × 错误："他对食材的要求极高，每一只虾都要亲自挑选。"（解说词，不是口播）
  × 错误："看，主理人在认真检查食材。"（纪录片旁白）
  口播台词 ≠ 旁白 ≠ 解说词。主理人就站在镜头前，自己开口说话。

■ R11【镜头结构防克隆 — 原创感的核心信号】
  · 每条脚本镜头数在5-8个之间，同一批3条的镜头数必须各不相同
  · 每条第1个镜头（开场）场景描述必须完全不同，禁止两条共用同一个开场场景
  · 同一批3条里，镜头顺序/组合禁止相同
  · 拍摄建议每条必须针对该条内容具体写，禁止复制粘贴同一句建议
  （镜头数不同 = 看起来像人手写，不像AI批量生产）

■ R12【去AI味 — 这是90分和70分的分水岭】
  程序会统计以下词在30条中的出现频次，超出限制的自动替换：
  · 每一个/每一件/每一位/每一道 → 全批超过3次后替换
  · 匠心 → 超过2次替换  · 用心/追求/坚持 → 超过3次替换
  · 初心/情怀/致力于/赋能 → 超过1-2次替换

  AI味的本质是：所有句子等长 + 概念堆砌 + 无具体人事物。
  消除AI味的方法：
  ① 句子一定要有长有短，最短可以只有2个字（"不行。""真的。""头大。"）
  ② 每条必须有一个具体的人/事/时间/金额，不能只有概念
  ③ 用"那个""那会儿""说实话""后来""其实"这类口语连接词
  ④ 允许不完整的句子，允许省略号停顿，允许破折号转折
  ⑤ 不要每段都有"结论句"，真人说话经常说了没有结论
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

JSON_FORMAT = """
请输出JSON格式（直接输出JSON数组，不要其他文字，不要markdown代码块）：
[
  {
    "number": 1,
    "type": "类型",
    "duration": "约1分钟",
    "title": "标题",
    "shots": [
      {"scene": "场景描述", "dialogue": "口播台词"},
      {"scene": "场景描述", "dialogue": "口播台词"}
    ],
    "tips": "拍摄建议"
  }
]
每条6-8个镜头，口播约1分钟，30条角度全部不同。
"""

# ── 自动重写：开场套话 + 违禁词（批量一次调用）────────
def rewrite_violations(scripts, violation_numbers, system_prompt, api_key, client_name):
    """对有开场套话或违禁词的脚本，批量重写违规部分"""
    if not violation_numbers:
        return scripts
    violating = [s for s in scripts if s.get('number') in violation_numbers]
    if not violating:
        return scripts

    api_client = openai.OpenAI(api_key=api_key, base_url="https://api.moonshot.cn/v1")
    rewrite_msg = f"""以下脚本有套话开场或违禁词，请根据脚本本身的内容进行修正。

主理人称呼：{client_name}

需要修正的脚本：
{json.dumps(violating, ensure_ascii=False, indent=2)}

修正要求：
1. 开场套话修正：只修改每条脚本 shots[0].dialogue（第一个镜头的口播）
   - 新的第一句必须根据该条脚本的具体内容来写，直接切入事件/数字/冲突/悬念
   - 禁止大家好/我是XX/欢迎/今天给大家等开头
   - 禁止使用"[填入]""[具体X]"这类占位符，必须写出实际内容
   - 字数控制在20字以内
   - 参考这条脚本的标题和其他镜头内容来创作开场句

2. 违禁词修正：把违禁词替换成自然表达，不是简单删除，要让语义完整流畅
   - 极限词（唯一/绝对/行业第一/史上最强等）→ 换成具体描述，比如"行业第一"改为"在这个圈子里做了十几年"
   - 效果类表述（快速见效/立竿见影/根治/治愈等）→ 改为真实感受描述，例如：
     "效果来得特别快的，一定要警惕成分是否温和。我的底线是，不添加任何违规刺激物。"
   - 如果是美业/护肤类脚本中提到具体效果改善的，在该句末尾自然加入：
     "这是客户的亲身感受，具体效果因人而异哦。"
   - 替换后的表达要保持口语自然，不能变成书面语

3. 其他所有字段保持不变
4. 输出格式：JSON数组，包含所有修正后的脚本（字段与输入完全一致）
直接输出JSON数组，不要其他文字，不要markdown代码块。"""

    full_text = ""
    try:
        stream = api_client.chat.completions.create(
            model="moonshot-v1-128k",
            max_tokens=8000,
            temperature=0.7,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": rewrite_msg},
            ],
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content
            if delta:
                full_text += delta
    except Exception:
        return scripts  # 重写失败则保留原文

    rewritten = extract_json_list(full_text)
    if rewritten:
        rewritten_map = {s.get('number'): s for s in rewritten}
        return [rewritten_map.get(s['number'], s) for s in scripts]
    return scripts

# ── 行业识别 ─────────────────────────────────────

EXTRA_KEYWORDS = {
    # 餐饮 — 具体菜品/食材/场景
    '牛排': 'canyin.md', '意面': 'canyin.md', '汉堡': 'canyin.md',
    '炸鸡': 'canyin.md', '披萨': 'canyin.md', '寿司': 'canyin.md',
    '甜品': 'canyin.md', '蛋糕': 'canyin.md', '炒饭': 'canyin.md',
    '龙虾': 'canyin.md', '皮皮虾': 'canyin.md', '海鲜': 'canyin.md',
    '螃蟹': 'canyin.md', '烤肉': 'canyin.md', '烤鱼': 'canyin.md',
    '串串': 'canyin.md', '卤味': 'canyin.md', '卤肉': 'canyin.md',
    '熟食': 'canyin.md', '江湖菜': 'canyin.md', '家常菜': 'canyin.md',
    '炒菜': 'canyin.md', '米粉': 'canyin.md', '粉面': 'canyin.md',
    '饺子': 'canyin.md', '包子': 'canyin.md', '馒头': 'canyin.md',
    '日料': 'canyin.md', '中餐': 'canyin.md', '西餐': 'canyin.md',
    '菜馆': 'canyin.md', '大排档': 'canyin.md', '夜宵': 'canyin.md',
    '螺蛳粉': 'canyin.md', '猪脚': 'canyin.md', '糕点': 'canyin.md',
    # 美业
    '剪发': 'meiiye.md', '发廊': 'meiiye.md', '护肤': 'meiiye.md',
    # 家装
    '装修': 'jiazhuang.md', '设计': 'jiazhuang.md',
}

INDUSTRY_NAMES = {
    'jiazhuang.md':  '家装',
    'canyin.md':     '餐饮',
    'meiiye.md':     '美业',
    'jiudian.md':    '酒店/宴席',
    'jinrong.md':    '金融/助贷',
    'gongchang.md':  '工厂/制造',
    'general.md':    '通用',
}

_VALID_SCRIPT_TYPES = {
    '故事型', '产品型', '幕后型', '痛点型', '教育型', '观点型', '对比型', '互动型', '氛围型',
    '创始人/故事类', '商业模式/加盟类', '行业分析类', '门店展示类', '痛点/避坑类', '招募/互动类',
    '其他',
}

def detect_industry(text):
    """
    计分式行业识别：统计每个行业命中的关键词数，得分最高的行业胜出。
    避免单个偶发词（如"装修"出现在"公积金装修贷"中）导致误判。
    """
    router = json.loads((PROMPTS_DIR / 'industry_router.json').read_text())
    all_rules = {**EXTRA_KEYWORDS, **{k: v for k, v in router['routing_rules'].items() if k != 'default'}}
    scores = {}
    for keyword, prompt_file in all_rules.items():
        if keyword in text:
            scores[prompt_file] = scores.get(prompt_file, 0) + 1
    if scores:
        return max(scores, key=lambda k: scores[k])
    return router['routing_rules']['default']

# ── 表单解析 ─────────────────────────────────────

def parse_form(raw):
    fields = {}
    current_key = None
    current_val = []
    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r'^[【⭐\*]*([^：:\n【】⭐\*]{2,20})[】]?[：:]\s*(.*)', line)
        if m:
            if current_key:
                fields[current_key] = '\n'.join(current_val).strip()
            current_key = m.group(1).strip()
            current_val = [m.group(2).strip()] if m.group(2).strip() else []
        elif current_key:
            current_val.append(line)
    if current_key:
        fields[current_key] = '\n'.join(current_val).strip()
    return fields

def get_field(fields, *keys, default=''):
    for k in keys:
        for fk, fv in fields.items():
            if k in fk and fv:
                return fv
    return default

def build_client(fields):
    name        = get_field(fields, '出镜称呼', '主理人姓名', '姓名')
    age         = get_field(fields, '年龄')
    gender      = get_field(fields, '性别')
    years       = get_field(fields, '从业年限', '年限')
    location    = extract_city(get_field(fields, '城市名字', '地点', '城市'))
    shop_pos    = extract_city(get_field(fields, '店铺位置', '位置'))
    shop        = get_field(fields, '店铺信息名称', '店名', '品牌名称', '公司名')
    hours       = get_field(fields, '营业时间')
    story       = get_field(fields, '创业经历', '入行故事', '故事')
    hard_time   = get_field(fields, '最难的时期', '最难熬', '最难', '困难', '低谷')
    best_case   = get_field(fields, '客户案例', '印象最深', '印象', '客户故事')
    diff        = get_field(fields, '与同行差异', '与同行', '差异', '不同')
    main_biz    = get_field(fields, '主营业务', '主营')
    product     = get_field(fields, '主推产品', '主推', '主营产品', '招牌产品')
    feature     = get_field(fields, '产品特点', '卖点', '特点')
    target      = get_field(fields, '受众人群', '目标客户', '客户群体', '目标客群')
    advantage   = get_field(fields, '核心优势', '优势')
    pain        = get_field(fields, '痛点', '解决客户')
    best        = get_field(fields, '客户喜好', '卖得多', '爆款')
    scale       = get_field(fields, '公司规模', '规模')
    b_or_c      = get_field(fields, 'B端', '对接')
    identity    = get_field(fields, '身份', '想体现')

    is_b2b = 'B' in b_or_c.upper() and 'C' not in b_or_c.upper()

    # 用所有字段值做识别，防止关键词出现在非预期字段里
    detect_text = " ".join(str(v) for v in fields.values() if v)
    prompt_file = detect_industry(detect_text)
    industry_keys = _get_industry_keys(detect_text)

    msg = f"""以下是客户信息，请生成30条短视频分镜脚本。

【出镜称呼】：{name}
【性别】：{gender}  【年龄】：{age}  【从业年限】：{years}
【店铺/品牌】：{shop}
【地点】：{location} {shop_pos}
【营业时间】：{hours}
【规模】：{scale}
【主营业务】：{main_biz}
【主推产品】：{product}
【产品特点/卖点】：{feature}
【创业经历/故事】：{story}
【最难熬的时期】：{hard_time}
【印象最深的客户案例】：{best_case}
【与同行最大的不同】：{diff}
【核心优势】：{advantage}
【目标受众】：{target}
【能解决的痛点】：{pain}
【热销产品】：{best}
【人设定位】：{identity}
【内容定位】：{"B端（招募合伙人/加盟商）" if is_b2b else "C端（面向消费者）"}

"""
    if is_b2b:
        msg += "内容角度覆盖：创始人故事、商业模式/加盟、行业分析、门店展示、痛点避坑、招募互动（30条不重复）\n"
        type_hint = "type只能是：创始人/故事类、商业模式/加盟类、行业分析类、门店展示类、痛点/避坑类、招募/互动类"
    else:
        msg += "内容形态必须多样，覆盖以下类型：故事型（个人经历/入行故事）、产品型（核心服务卖点）、幕后型（工作流程/日常）、痛点型（行业避坑/内幕）、教育型（干货知识/行业科普）、观点型（立场/主理人独特见解）、对比型（前后变化/真实案例）、互动型（问答/投票/粉丝共创）、氛围型（场景日常/情感共鸣）。30条角度完全不同，同一类型不超过6条。\n"
        msg += "分布：第1条和第2条必须是故事型（人设介绍），故事型4-6条，产品型3-5条，幕后型3-4条，痛点型4-5条，教育型3-4条，观点型2-3条，对比型2-3条，互动型2-3条，氛围型2-3条。\n"
        type_hint = "type只能是：故事型、产品型、幕后型、痛点型、教育型、观点型、对比型、互动型、氛围型"

    msg += "\n" + type_hint

    company = shop or main_biz or name

    client = {
        'id':            re.sub(r'\W', '', name)[:8],
        'name':          name,
        'company':       company,
        'location':      f"{location} {shop_pos}".strip(),
        'prompt_file':   prompt_file,
        'industry_keys': industry_keys,
        'user_message':  msg,
    }
    return client, prompt_file

# ── API 调用 ──────────────────────────────────────

def extract_json_list(text):
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    match = re.search(r'\[\s*\{[\s\S]*\}\s*\]', text)
    if not match:
        match = re.search(r'\[[\s\S]+\]', text)
    if not match:
        return []
    try:
        return json.loads(match.group())
    except json.JSONDecodeError:
        return []

def generate_scripts(client, api_key, progress_callback=None):
    """
    v2.1 三层防崩架构：大纲 → 10批×3条（每批重注入规则）→ 违禁词扫描
    每批只生成3条，防止模型在同批次内耗尽创意角度导致末尾脚本重复。
    progress_callback(pct, message) 用于向UI汇报进度，可选。
    """
    api_client = openai.OpenAI(api_key=api_key, base_url="https://api.moonshot.cn/v1")
    system_prompt = load_prompt(client['prompt_file'])
    base_msg = client['user_message']

    def kimi_call(user_msg, max_tokens=8000):
        full_text = ""
        stream = api_client.chat.completions.create(
            model="moonshot-v1-128k",
            max_tokens=max_tokens,
            temperature=0.7,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_msg},
            ],
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content
            if delta:
                full_text += delta
        return full_text

    # ── 阶段一：生成30条选题大纲 ─────────────────────
    if progress_callback:
        progress_callback(0.05, "阶段一：正在生成30条选题大纲...")

    outline_msg = base_msg + """

生成30条选题大纲，严格遵守：
- 第1条和第2条必须是故事型（人设介绍，写清楚讲哪个具体经历）
- 内容形态要多样：故事型4-6条、产品型3-5条、幕后型3-4条、痛点型4-5条、教育型3-4条、观点型2-3条、对比型2-3条、互动型2-3条、氛围型2-3条，同一类型不超过6条
- 30条核心角度完全不同，不得有两条讲同一件具体事件
- 标题要说清楚讲什么具体事件或内容，不能是泛泛概念（"我的故事"这种标题不合格）
- 教育型：直接给一个行业干货知识点；观点型：主理人对某件事的独特看法；对比型：具体的前后变化或案例
输出格式（JSON数组，不要其他文字）：
[{"number":1,"type":"故事型","title":"说清楚讲哪个具体事件的标题"},...]"""

    raw_outline = kimi_call(outline_msg, max_tokens=3000)
    outline = extract_json_list(raw_outline)

    outline_text = ""
    if outline:
        outline_text = "\n\n【已确定30条选题大纲】\n" + "\n".join(
            f"{o.get('number', i+1):02d}. [{o.get('type','')}] {o.get('title','')}"
            for i, o in enumerate(outline)
        )

    # ── 阶段二：10批×3条，每批重注入规则包 ────────────
    # 每批只生成3条：防止模型在同批次内创意耗尽，导致第4-5条重复前几条
    all_scripts = []
    failed_batches = []

    for batch_idx in range(10):
        batch_start = batch_idx * 3 + 1
        batch_end   = batch_start + 2
        progress    = 0.1 + batch_idx * 0.08

        if progress_callback:
            progress_callback(progress, f"阶段二：正在生成第{batch_start}–{batch_end}条（批次{batch_idx+1}/10）...")

        # 防重复摘要：标题+开场台词+开场场景，三维防克隆
        anti_repeat = ""
        if all_scripts:
            anti_repeat = "\n\n【防重复 — 以下条目已生成，本批禁止重复其核心事件、开场句式、开场镜头场景】\n" + "\n".join(
                "第{}条: [{}] {} | 开场台词: {} | 开场场景: {}".format(
                    s['number'], s.get('type', ''),
                    s.get('title', ''),
                    (s.get('shots') or [{}])[0].get('dialogue', '')[:20],
                    (s.get('shots') or [{}])[0].get('scene', '')[:15],
                )
                for s in all_scripts
            )

        # 本批大纲条目
        batch_outline = ""
        if outline:
            batch_items = outline[batch_idx * 3: (batch_idx + 1) * 3]
            batch_outline = "\n\n【本批大纲（第{}-{}条）】\n".format(batch_start, batch_end) + "\n".join(
                "{:02d}. [{}] {}".format(
                    o.get('number', batch_start + i), o.get('type', ''), o.get('title', '')
                )
                for i, o in enumerate(batch_items)
            )

        # 行业专项违禁词注入
        industry_forbidden = ""
        ikeys = client.get('industry_keys', [])
        if ikeys:
            extra_words = []
            for k in ikeys:
                extra_words += _FW_DATA.get(k, [])
            if extra_words:
                industry_forbidden = "\n\n【本行业专项违禁词 — 以下词绝对不能出现在口播或场景描述中】\n" + "、".join(extra_words[:40])

        batch_msg = (
            base_msg
            + outline_text
            + anti_repeat
            + RULES_V2          # 每批重注入完整规则包
            + industry_forbidden
            + batch_outline
            + f"""

现在生成第{batch_start}条到第{batch_end}条完整分镜脚本（共3条）。

【镜头数与字数要求 — 最重要】
每条脚本的镜头数在5-8个之间，这3条的镜头数必须各不相同（比如5、7、6，禁止3条都一样）。
每个镜头口播30-40字，全条合计必须达到150字以上（上限210字）。
· 5镜头：每镜头平均32字 ≈ 160字
· 6镜头：每镜头平均28字 ≈ 170字
· 7镜头：每镜头平均25字 ≈ 175字
· 8镜头：每镜头平均22字 ≈ 176字
禁止每个镜头只写一两句就结束，口播要说完整、有情绪、有细节。

这3条必须角度完全不同，开场句式完全不同，镜头结构完全不同。
直接输出JSON数组，不要其他文字，不要markdown代码块：
[
  {{
    "number": {batch_start},
    "type": "类型",
    "duration": "约1分钟",
    "title": "标题",
    "shots": [
      {{"scene": "场景描述10-20字", "dialogue": "口播台词30-40字，有具体细节和情绪"}},
      {{"scene": "场景描述10-20字", "dialogue": "口播台词30-40字"}},
      {{"scene": "场景描述10-20字", "dialogue": "口播台词30-40字"}},
      {{"scene": "场景描述10-20字", "dialogue": "口播台词30-40字"}},
      {{"scene": "场景描述10-20字", "dialogue": "口播台词30-40字，结尾按类型处理"}}
    ],
    "tips": "拍摄建议（门店内/手机可拍，一句话，针对本条内容具体写）"
  }},
  ...共3条，每条镜头数不同
]"""
        )

        raw = kimi_call(batch_msg, max_tokens=8000)
        parsed = extract_json_list(raw)
        if parsed:
            for i, s in enumerate(parsed):
                s['number'] = batch_start + i
                if 'difficulty' not in s:
                    s['difficulty'] = '深度版'
                if s.get('type') not in _VALID_SCRIPT_TYPES:
                    s['type'] = '其他'
            all_scripts.extend(parsed)
        else:
            failed_batches.append(f"{batch_start}-{batch_end}")

    # ── 阶段三：程序级扫描 ─────────────────────────────
    opening_violations = scan_forbidden_openings(all_scripts)
    word_violations    = scan_forbidden_words(all_scripts, client['prompt_file'], client.get('industry_keys', []))
    scene_violations   = scan_scene_persons(all_scripts)

    # 合并需要重写的条号
    rewrite_numbers = set(opening_violations) | {v['script'] for v in word_violations}

    # 触发自动重写（开场套话 + 违禁词）
    if rewrite_numbers:
        if progress_callback:
            progress_callback(0.88, f"⚠️ 发现{len(rewrite_numbers)}条违规，正在自动重写...")
        all_scripts = rewrite_violations(
            all_scripts, rewrite_numbers,
            system_prompt, api_key, client['name']
        )
        # 重写后重新扫描一次
        opening_violations = scan_forbidden_openings(all_scripts)
        word_violations    = scan_forbidden_words(all_scripts, client['prompt_file'], client.get('industry_keys', []))

    # ── 阶段四：程序级后处理（极限词+词频+C2/C4+结尾配额）──
    all_scripts, zui_replacements = post_process_scripts(all_scripts, client.get('name', ''))

    if failed_batches:
        client['failed_batches'] = failed_batches

    # ── 阶段五：自动评分（仅展示，不触发重试）─────────────
    if progress_callback:
        progress_callback(0.96, "正在评分...")
    score_result = auto_score(all_scripts, client)
    client['score'] = score_result

    if progress_callback:
        progress_callback(0.98, f"生成完成，共{len(all_scripts)}条，正在制作Word...")

    client['violations']                    = word_violations
    client['scene_violations']              = scene_violations
    client['opening_violations_remaining']  = opening_violations
    client['zui_replacements']              = zui_replacements
    return all_scripts

# ── Word 生成（返回字节流，适配Web下载）────────────

def make_word_bytes(client, scripts):
    doc = Document()

    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2.2)
        sec.top_margin  = sec.bottom_margin = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(client['company'])
    r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('30条短视频分镜脚本')
    r2.font.size = Pt(18); r2.font.bold = True

    doc.add_paragraph()
    for line in [
        f'主理人：{client["name"]}',
        f'地点：{client["location"]}',
        f'制作日期：{datetime.date.today().strftime("%Y年%m月%d日")}',
        '晓牧传媒 出品',
    ]:
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ri = pi.add_run(line)
        ri.font.size = Pt(11)
        ri.font.color.rgb = RGBColor(0x66,0x66,0x66)
        pi.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    easy_n = sum(1 for s in scripts if s.get('difficulty')=='简单版')
    std_n  = sum(1 for s in scripts if s.get('difficulty')=='标准版')
    deep_n = sum(1 for s in scripts if s.get('difficulty')=='深度版')
    for i, (label, bg, desc) in enumerate([
        (f'🟢 简单版 {easy_n}条', 'E8F5E9', '3-4镜头 · 20-30秒'),
        (f'🟠 标准版 {std_n}条',  'FFF3E0', '5镜头 · 30-60秒'),
        (f'🟣 深度版 {deep_n}条', 'EDE7F6', '6-8镜头 · 约1分钟'),
    ]):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, bg)
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pp.add_run(f'{label}\n'); rl.font.bold = True; rl.font.size = Pt(10)
        rd = pp.add_run(desc); rd.font.size = Pt(9)
        rd.font.color.rgb = RGBColor(0x55,0x55,0x55)

    doc.add_page_break()

    # 按编号顺序 1–30 输出，类型信息已在每条脚本标签栏中显示
    for s in sorted(scripts, key=lambda s: s.get('number', 999)):
        _add_script_block(doc, s)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_script_block(doc, s):
    fg_t, bg_hex = difficulty_style(s.get('difficulty','深度版'))
    fg_color = RGBColor(*fg_t)
    shots = s.get('shots', [])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r_n = p.add_run(f'第{s.get("number","")}条  ')
    r_n.font.size = Pt(13); r_n.font.bold = True
    r_n.font.color.rgb = RGBColor(0xCC,0x55,0x00)
    r_t = p.add_run(s.get('title',''))
    r_t.font.size = Pt(13); r_t.font.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    for tag, color in [
        (f'  {s.get("type","")}  ',       RGBColor(0x55,0x55,0x55)),
        (f'  {s.get("difficulty","")}  ',  fg_color),
        (f'  ⏱ {s.get("duration","")}  ', RGBColor(0x77,0x77,0x77)),
        (f'  📷 {len(shots)}个镜头  ',     RGBColor(0x77,0x77,0x77)),
    ]:
        r = p2.add_run(tag); r.font.size = Pt(9); r.font.color.rgb = color

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdrs = table.rows[0].cells
    for cell, txt in zip(hdrs, ['镜头', '场景描述', '口播内容']):
        set_cell_bg(cell, '1A1A2E')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(txt)
        r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    for i, shot in enumerate(shots):
        row = table.add_row()
        c0 = row.cells[0]
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = c0.paragraphs[0].add_run(f'镜头{i+1}')
        r0.font.size = Pt(10); r0.font.bold = True
        c1 = row.cells[1]
        r1 = c1.paragraphs[0].add_run(shot.get('scene',''))
        r1.font.size = Pt(10); r1.font.italic = True
        r1.font.color.rgb = RGBColor(0x44,0x44,0x44)
        c2 = row.cells[2]
        r2 = c2.paragraphs[0].add_run(shot.get('dialogue',''))
        r2.font.size = Pt(10.5)
        if i % 2 == 1:
            for c in [c0, c1, c2]:
                set_cell_bg(c, 'F9F9F9')

    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(7.5)

    p_tip = doc.add_paragraph()
    p_tip.paragraph_format.space_before = Pt(5)
    p_tip.paragraph_format.space_after  = Pt(12)
    rl = p_tip.add_run('📌 拍摄建议  ')
    rl.font.size = Pt(10); rl.font.bold = True
    rl.font.color.rgb = RGBColor(0xC0,0x62,0x00)
    rt = p_tip.add_run(s.get('tips',''))
    rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0x44,0x44,0x44)

    pd = doc.add_paragraph()
    pd.add_run('─' * 58).font.size = Pt(7)
    pd.runs[0].font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    pd.paragraph_format.space_after = Pt(2)
